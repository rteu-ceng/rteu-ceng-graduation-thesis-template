"""
RTEÜ Bilgisayar Mühendisliği - Bitirme Tezi  Şablon Oluşturucu
Tüm Özellikler Birleştirilmiş Versiyon
Author: Dr. Uğur CORUH
Version: 5.0 -  Edition
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class RTEUThesisTemplate:
    """RTEÜ Bilgisayar Mühendisliği  Bitirme Tezi Şablonu"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.define_styles()
        
    def setup_document(self):
        """Sayfa yapılandırması"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.page_width = Cm(21)  # A4
            section.page_height = Cm(29.7)  # A4
            
    def define_styles(self):
        """Özel stil tanımlamaları"""
        styles = self.doc.styles
        
        # Başlık 1
        heading1 = styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 128)
        
        # Başlık 2
        heading2 = styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 0, 0)
        
        # Başlık 3
        heading3 = styles['Heading 3']
        heading3.font.name = 'Calibri'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.font.color.rgb = RGBColor(64, 64, 64)
        
        # Normal metin
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        
    def add_header_footer(self):
        """Üst ve alt bilgi ekle"""
        # Üst bilgi
        header = self.doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "RTEÜ Bilgisayar Mühendisliği - Bitirme Tezi 2025-2026"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.style.font.size = Pt(10)
        header_para.style.font.italic = True
        
        # Alt bilgi
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Oluşturma Tarihi: {datetime.now().strftime('%d.%m.%Y')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_para.style.font.size = Pt(9)
        
    def add_checkbox(self, paragraph, text):
        """Checkbox ekle"""
        run = paragraph.add_run()
        run.add_text('☐ ')
        run.font.size = Pt(11)
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
        
    def add_fillable_field(self, paragraph, label, width=50):
        """Doldurulabilir alan ekle"""
        run = paragraph.add_run(f"{label}: ")
        run.font.bold = True
        run = paragraph.add_run("_" * width)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
    def create_cover_page(self):
        """Kapak sayfası oluştur"""
        # Logo ve başlık
        title = self.doc.add_heading('', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('RECEP TAYYİP ERDOĞAN ÜNİVERSİTESİ\n')
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)
        
        run = title.add_run('Mühendislik ve Mimarlık Fakültesi\n')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        run = title.add_run('Bilgisayar Mühendisliği Bölümü')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Ana başlık
        main_title = self.doc.add_heading('', level=1)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = main_title.add_run('🚀 BİTİRME TEZİ FİKİR ÖNERİSİ')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('📚 2025-2026 Akademik Yılı\n')
        run.font.size = Pt(14)
        run = subtitle.add_run('💡 Teknoloji + Girişimcilik = Unicorn')
        run.font.size = Pt(14)
        run.font.italic = True
        
        # Motivasyon sözü
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        quote = self.doc.add_paragraph()
        quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        quote.add_run('"The best way to predict the future is to invent it."\n')
        quote.add_run('- Alan Kay')
        quote.runs[0].font.size = Pt(14)
        quote.runs[0].font.italic = True
        quote.runs[1].font.size = Pt(12)
        
        self.doc.add_page_break()
        
    def create_team_info_section(self):
        """Takım bilgileri bölümü"""
        self.doc.add_heading('📋 TAKIM BİLGİLERİ', level=1)
        
        # Takım adı ve proje bilgileri
        p = self.doc.add_paragraph()
        self.add_fillable_field(p, "Takım Adı", 60)
        
        p = self.doc.add_paragraph()
        self.add_fillable_field(p, "Proje Başlığı", 80)
        
        p = self.doc.add_paragraph()
        self.add_fillable_field(p, "Fikir No", 10)
        p.add_run(" / 5")
        
        p = self.doc.add_paragraph()
        self.add_fillable_field(p, "Proje Sloganı", 70)
        
        # Takım üyeleri tablosu
        self.doc.add_heading('Takım Üyeleri', level=2)
        table = self.doc.add_table(rows=4, cols=7)
        table.style = 'Table Grid'
        
        # Başlık satırı
        headers = ['Rol', 'Ad Soyad', 'Öğrenci No', 'E-posta', 'GitHub', 'LinkedIn', 'İletişim']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Takım rolleri
        roles = ['Takım Lideri', 'Teknik Lider', 'Üye 3 (Opsiyonel)']
        for i, role in enumerate(roles, 1):
            table.cell(i, 0).text = role
            table.cell(i, 3).text = '@erdogan.edu.tr'
            
        # Takım sözleşmesi
        self.doc.add_heading('Takım Sözleşmesi', level=2)
        contract_items = [
            'Haftalık düzenli toplantı yapacağız',
            'GitHub\'a düzenli commit atacağız (min. 5/hafta)',
            'Kod review yapacağız',
            'Dokümantasyonu güncel tutacağız',
            'İletişim kanallarını aktif kullanacağız',
            'Görev dağılımına uyacağız',
            'Deadline\'lara sadık kalacağız'
        ]
        
        for item in contract_items:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, item)
            
    def create_comprehensive_domains(self):
        """Genişletilmiş proje domain kategorileri"""
        self.doc.add_page_break()
        self.doc.add_heading('🎯 PROJE DOMAIN KATEGORİLERİ', level=1)
        
        self.doc.add_paragraph('Lütfen projenizin ana domain alanını ve alt kategorilerini seçiniz. ' + 
                              'Birden fazla domain seçilebilir. Listede olmayan alanları "Diğer" kısmına yazınız.')
        
        # Tüm domain kategorileri - Genişletilmiş liste
        all_domains = {
            '🤖 Yapay Zeka & Makine Öğrenmesi': [
                'Computer Vision & Image Processing',
                'Natural Language Processing (NLP)',
                'Speech Recognition & Synthesis',
                'Reinforcement Learning',
                'Generative AI (LLM, Diffusion Models)',
                'Explainable AI (XAI)',
                'TinyML & Edge AI',
                'AutoML & Neural Architecture Search',
                'Federated Learning',
                'Quantum Machine Learning',
                'AI Ethics & Fairness',
                'Multimodal AI',
                'Few-shot & Zero-shot Learning',
                'Graph Neural Networks',
                'Neuro-symbolic AI',
                'AI in Healthcare Diagnostics',
                'Predictive Maintenance AI',
                'Conversational AI & Chatbots',
                'Emotion Recognition',
                'AI-powered Recommendation Systems'
            ],
            
            '🔒 Siber Güvenlik & Gizlilik': [
                'Zero Trust Architecture',
                'Blockchain Security',
                'Post-Quantum Cryptography',
                'SIEM & SOC Automation',
                'Privacy-Preserving ML',
                'IoT Security',
                'Cloud Security',
                'Application Security (SAST/DAST)',
                'Network Security & Firewall',
                'Identity & Access Management',
                'Threat Intelligence',
                'Digital Forensics',
                'Malware Analysis',
                'Security Orchestration (SOAR)',
                'DevSecOps',
                'Hardware Security',
                'Mobile Security',
                'API Security',
                'Container Security',
                'Ransomware Protection'
            ],
            
            '📊 Veri Bilimi & Büyük Veri': [
                'Big Data Processing (Spark, Hadoop)',
                'Real-time Analytics & Stream Processing',
                'Business Intelligence & Reporting',
                'Predictive Analytics',
                'Data Engineering & ETL/ELT',
                'DataOps & MLOps',
                'Data Governance & Quality',
                'Time Series Analysis',
                'Graph Analytics',
                'Data Warehousing',
                'Data Lakes & Lakehouses',
                'Feature Engineering',
                'A/B Testing & Experimentation',
                'Data Visualization & Dashboards',
                'Geospatial Analytics',
                'Data Mesh Architecture',
                'Synthetic Data Generation',
                'Data Privacy & Anonymization',
                'Apache Kafka & Event Streaming',
                'Data Science Platforms'
            ],
            
            '🌐 Web3 & Blockchain': [
                'DeFi (Decentralized Finance)',
                'Smart Contracts Development',
                'NFTs & Digital Assets',
                'Cross-chain Solutions & Bridges',
                'Layer 2 Solutions',
                'Decentralized Storage (IPFS, Filecoin)',
                'DAO & Governance',
                'DApps Development',
                'Blockchain Oracles',
                'Tokenomics & Cryptoeconomics',
                'Consensus Mechanisms',
                'Web3 Gaming & Metaverse',
                'Decentralized Identity (DID)',
                'Supply Chain on Blockchain',
                'CBDC & Digital Currencies',
                'Zero-Knowledge Proofs',
                'Blockchain Interoperability',
                'DeFi Protocols',
                'Blockchain Analytics',
                'Crypto Wallets'
            ],
            
            '☁️ Bulut Bilişim & DevOps': [
                'Kubernetes & Container Orchestration',
                'Serverless & FaaS',
                'Infrastructure as Code (Terraform, Ansible)',
                'Multi-cloud & Hybrid Cloud',
                'Site Reliability Engineering (SRE)',
                'GitOps & CI/CD',
                'Service Mesh (Istio, Linkerd)',
                'Cloud Native Development',
                'Observability & Monitoring',
                'Chaos Engineering',
                'Platform Engineering',
                'Edge Computing & CDN',
                'Cloud Cost Optimization',
                'Backup & Disaster Recovery',
                'Cloud Migration Strategies',
                'FinOps',
                'Cloud Security Posture Management',
                'Container Registry Management',
                'Blue-Green Deployments',
                'Infrastructure Monitoring'
            ],
            
            '📱 Mobil & Cross-Platform': [
                'Native iOS Development (Swift)',
                'Native Android Development (Kotlin)',
                'React Native Development',
                'Flutter Development',
                'Progressive Web Apps (PWA)',
                'Hybrid App Development',
                'Mobile Game Development',
                'AR/VR Mobile Apps',
                'Mobile Payment Systems',
                'Mobile Health Apps',
                'Mobile DevOps',
                'App Store Optimization',
                'Mobile Analytics',
                'Wearable App Development',
                'Mobile Security & Privacy',
                'Mobile UI/UX Design',
                'Push Notification Services',
                'Mobile Backend as a Service',
                'Mobile Testing Automation',
                'Offline-First Mobile Apps'
            ],
            
            '🌍 Web Teknolojileri': [
                'Full-Stack Development',
                'JAMstack Architecture',
                'Micro-frontend Architecture',
                'Real-time Web (WebRTC, WebSockets)',
                'GraphQL & REST APIs',
                'Server-Side Rendering (SSR)',
                'Static Site Generation (SSG)',
                'Web Performance Optimization',
                'Web Accessibility (WCAG)',
                'Web Components',
                'WebAssembly',
                'Browser Extensions',
                'E-commerce Platforms',
                'Content Management Systems',
                'Web Analytics & SEO',
                'Progressive Enhancement',
                'Single Page Applications',
                'Web Security Best Practices',
                'API Gateway Management',
                'Headless CMS'
            ],
            
            '🔌 IoT & Gömülü Sistemler': [
                'Industrial IoT (IIoT)',
                'Smart City Solutions',
                'Smart Home Automation',
                'Wearable Technology',
                'Edge Computing Devices',
                'Digital Twins',
                'Sensor Networks & WSN',
                'RTOS Development',
                'Embedded Linux',
                'FPGA Programming',
                'Arduino & Raspberry Pi',
                'LoRaWAN & NB-IoT',
                'MQTT & CoAP Protocols',
                'Energy Harvesting',
                'Robotics & Drones',
                'IoT Platforms',
                'Device Management',
                'IoT Analytics',
                'Connected Vehicles',
                'Agricultural IoT'
            ],
            
            '🎮 Oyun & Eğlence Teknolojileri': [
                'Game Engine Development',
                'Unity 3D Development',
                'Unreal Engine Development',
                'Mobile Game Development',
                'Multiplayer & Networking',
                'Game AI & NPC Behavior',
                'Procedural Generation',
                'VR/AR Gaming',
                'Game Analytics',
                'Game Monetization',
                'Esports Platforms',
                'Cloud Gaming',
                'Game Streaming',
                'Serious Games & Gamification',
                'Game Physics & Simulation',
                'Level Design Tools',
                'Game Audio Programming',
                'Cross-platform Gaming',
                'Game Testing & QA',
                'In-game Economy Design'
            ],
            
            '🏥 Sağlık Teknolojileri': [
                'Telemedicine Platforms',
                'Medical Image Analysis',
                'Electronic Health Records (EHR)',
                'Health Information Systems',
                'Wearable Health Monitoring',
                'Drug Discovery & AI',
                'Mental Health Apps',
                'Remote Patient Monitoring',
                'Clinical Decision Support',
                'Bioinformatics',
                'Medical IoT Devices',
                'Health Data Analytics',
                'Digital Therapeutics',
                'Personalized Medicine',
                'Healthcare Blockchain',
                'Medical Robotics',
                'Healthcare Chatbots',
                'Fitness & Wellness Apps',
                'Healthcare Compliance',
                'Medical Device Software'
            ],
            
            '💰 FinTech & InsurTech': [
                'Digital Banking Solutions',
                'Payment Processing Systems',
                'Robo-advisors',
                'Cryptocurrency Exchanges',
                'Risk Assessment & Credit Scoring',
                'RegTech & Compliance',
                'Open Banking APIs',
                'Mobile Wallets',
                'P2P Lending Platforms',
                'Algorithmic Trading',
                'Fraud Detection Systems',
                'Insurance Automation',
                'KYC/AML Solutions',
                'Financial Planning Apps',
                'Blockchain in Finance',
                'Neo Banks',
                'Expense Management',
                'Investment Platforms',
                'Financial Data Aggregation',
                'Tax Technology'
            ],
            
            '🎓 EdTech & E-Learning': [
                'Learning Management Systems (LMS)',
                'Virtual Classrooms',
                'Educational Games',
                'Adaptive Learning Platforms',
                'Online Course Platforms',
                'Student Assessment Tools',
                'AR/VR in Education',
                'AI Tutoring Systems',
                'Plagiarism Detection',
                'Educational Analytics',
                'Skill Assessment Platforms',
                'Language Learning Apps',
                'STEM Education Tools',
                'Collaborative Learning Tools',
                'Microlearning Platforms',
                'Corporate Training Platforms',
                'Educational Content Creation',
                'Student Information Systems',
                'Proctoring Solutions',
                'Knowledge Management'
            ],
            
            '🚗 Otomotiv & Ulaşım': [
                'Autonomous Vehicle Systems',
                'Connected Car Technology',
                'Fleet Management Systems',
                'Traffic Management & Smart Cities',
                'Electric Vehicle Software',
                'Vehicle Telematics',
                'Ride-sharing Platforms',
                'Parking Solutions',
                'Navigation & Mapping',
                'Vehicle-to-Everything (V2X)',
                'Driver Assistance Systems',
                'Mobility as a Service (MaaS)',
                'Drone Delivery Systems',
                'Transportation Analytics',
                'Public Transit Solutions',
                'Logistics Optimization',
                'Last-Mile Delivery',
                'Vehicle Diagnostics',
                'Car Rental Platforms',
                'Traffic Prediction'
            ],
            
            '🏭 Endüstri 4.0 & Üretim': [
                'Manufacturing Execution Systems (MES)',
                'Predictive Maintenance',
                'Quality Control Automation',
                'Supply Chain Management',
                'Digital Twin Manufacturing',
                'Industrial Robotics',
                'Computer-Aided Manufacturing',
                'Production Planning & Scheduling',
                'Inventory Management',
                'Factory Automation',
                'Process Optimization',
                'Industrial AR/VR',
                'SCADA Systems',
                'Energy Management',
                'Lean Manufacturing Tools',
                'Asset Management',
                'Warehouse Automation',
                '3D Printing Software',
                'Process Mining',
                'OEE Monitoring'
            ],
            
            '🌱 Çevre & Sürdürülebilirlik': [
                'Carbon Footprint Tracking',
                'Renewable Energy Management',
                'Smart Grid Solutions',
                'Waste Management Systems',
                'Water Management',
                'Environmental Monitoring',
                'Sustainable Agriculture Tech',
                'Green Building Solutions',
                'Circular Economy Platforms',
                'Climate Change Analytics',
                'Energy Efficiency Tools',
                'Emission Monitoring',
                'Recycling Technology',
                'Biodiversity Tracking',
                'ESG Reporting Tools',
                'Clean Tech Solutions',
                'Pollution Monitoring',
                'Sustainable Supply Chain',
                'Carbon Trading Platforms',
                'Environmental Compliance'
            ],
            
            '🛍️ E-Ticaret & Perakende': [
                'E-commerce Platforms',
                'Marketplace Development',
                'Inventory Management',
                'Order Management Systems',
                'Customer Analytics',
                'Recommendation Engines',
                'Dynamic Pricing',
                'Shopping Cart Solutions',
                'Payment Gateway Integration',
                'Dropshipping Platforms',
                'Social Commerce',
                'Omnichannel Retail',
                'Loyalty Programs',
                'Product Information Management',
                'Returns Management',
                'Live Commerce',
                'B2B E-commerce',
                'Subscription Commerce',
                'Visual Commerce',
                'Voice Commerce'
            ],
            
            '🎨 Yaratıcı Teknolojiler': [
                'Generative Art & Design',
                'Music Generation AI',
                'Video Editing Automation',
                '3D Modeling & Animation',
                'Digital Art Platforms',
                'Content Creation Tools',
                'Deepfake Technology',
                'Voice Synthesis',
                'Photo Editing AI',
                'NFT Marketplaces',
                'Virtual Production',
                'Motion Capture',
                'Digital Fashion',
                'Augmented Reality Filters',
                'Creative Collaboration Tools',
                'AI Writing Assistants',
                'Podcast Production Tools',
                'Live Streaming Platforms',
                'Virtual Events Platforms',
                'Digital Asset Management'
            ],
            
            '🏛️ GovTech & Kamu Teknolojileri': [
                'E-Government Platforms',
                'Digital Identity Systems',
                'Online Voting Systems',
                'Public Service Automation',
                'Smart City Infrastructure',
                'Emergency Response Systems',
                'Tax Management Systems',
                'Document Management',
                'Citizen Engagement Platforms',
                'Government Data Portals',
                'Public Safety Solutions',
                'Court Management Systems',
                'License & Permit Systems',
                'Social Welfare Platforms',
                'Transparency & Anti-corruption Tools',
                'Public Health Systems',
                'Border Control Systems',
                'Municipal Services',
                'Legislative Management',
                'Public Procurement'
            ],
            
            '🚀 Uzay & Havacılık': [
                'Satellite Data Processing',
                'Ground Station Software',
                'Flight Management Systems',
                'Drone Control Systems',
                'Space Mission Planning',
                'Orbital Mechanics Software',
                'Remote Sensing Applications',
                'Aviation Safety Systems',
                'Air Traffic Management',
                'Spacecraft Simulation',
                'Rocket Telemetry',
                'Weather Prediction Systems',
                'Navigation Systems',
                'Space Debris Tracking',
                'Aerospace Testing Tools',
                'Mission Control Software',
                'Satellite Communication',
                'UAV Path Planning',
                'Aircraft Maintenance',
                'Space Tourism Platforms'
            ],
            
            '⚡ Enerji Teknolojileri': [
                'Smart Grid Management',
                'Energy Trading Platforms',
                'Solar Panel Optimization',
                'Wind Farm Management',
                'Battery Management Systems',
                'Energy Storage Solutions',
                'Demand Response Systems',
                'Power Plant Automation',
                'Energy Analytics',
                'Grid Stability Tools',
                'Microgrid Solutions',
                'EV Charging Infrastructure',
                'Energy Efficiency Monitoring',
                'Renewable Energy Forecasting',
                'Carbon Credit Platforms',
                'Energy Billing Systems',
                'Peak Load Management',
                'Energy Audit Tools',
                'Distributed Energy Resources',
                'Energy Market Analytics'
            ]
        }
        
        # Domain kategorilerini ekle
        for domain, subcategories in all_domains.items():
            p = self.doc.add_paragraph()
            self.add_checkbox(p, domain)
            p.runs[1].font.bold = True
            p.runs[1].font.size = Pt(12)
            p.runs[1].font.color.rgb = RGBColor(0, 0, 128)
            
            # Alt kategoriler (ilk 15'i göster)
            for subcat in subcategories[:15]:
                sub_p = self.doc.add_paragraph(style='List Bullet')
                sub_p.paragraph_format.left_indent = Inches(0.5)
                self.add_checkbox(sub_p, subcat)
            
            # Diğer seçeneği ve boş alan
            other_p = self.doc.add_paragraph(style='List Bullet')
            other_p.paragraph_format.left_indent = Inches(0.5)
            self.add_checkbox(other_p, 'Diğer (Belirtiniz): ')
            
            # Boş yazma alanı
            text_area = self.doc.add_paragraph()
            text_area.paragraph_format.left_indent = Inches(0.7)
            text_area.add_run('_' * 80 + '\n' + '_' * 80)
            text_area.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            self.doc.add_paragraph()  # Boşluk
            
    def create_comprehensive_tech_stack(self):
        """Genişletilmiş teknoloji stack'i"""
        self.doc.add_page_break()
        self.doc.add_heading('💻 TEKNOLOJİ STACK\'İ VE ARAÇLAR', level=1)
        
        self.doc.add_paragraph('Projenizde kullanmayı planladığınız teknolojileri işaretleyiniz. ' + 
                              'Listede olmayan teknolojileri "Diğer" kısmına ekleyiniz.')
        
        # Genişletilmiş teknoloji kategorileri
        tech_categories = {
            '🎨 Frontend Teknolojileri': {
                'JavaScript Frameworks': ['React 18+', 'Vue 3', 'Angular 17+', 'Svelte', 'SolidJS', 
                                         'Preact', 'Alpine.js', 'Lit', 'Qwik', 'Astro'],
                'Meta-Frameworks': ['Next.js 14+', 'Nuxt 3', 'Remix', 'Gatsby', 'Astro', 
                                   'SvelteKit', 'Qwik City', 'Fresh', 'T3 Stack', 'Blitz.js'],
                'CSS Frameworks': ['Tailwind CSS', 'Bootstrap 5', 'Material UI', 'Ant Design', 
                                  'Chakra UI', 'Bulma', 'Foundation', 'Semantic UI', 'Mantine', 'PrimeReact'],
                'Build Tools': ['Vite', 'Webpack', 'Parcel', 'esbuild', 'Rollup', 'SWC', 
                               'Turbopack', 'Bun', 'Rome', 'Nx'],
                'State Management': ['Redux Toolkit', 'Zustand', 'MobX', 'Recoil', 'Jotai', 
                                    'Valtio', 'XState', 'Pinia', 'Tanstack Query', 'SWR'],
                'Testing': ['Jest', 'Vitest', 'Cypress', 'Playwright', 'Testing Library', 
                           'Storybook', 'Puppeteer', 'WebDriver', 'Karma', 'Mocha']
            },
            
            '⚙️ Backend Teknolojileri': {
                'Node.js': ['Express.js', 'Fastify', 'NestJS', 'Koa', 'Hapi', 'AdonisJS', 
                           'Feathers', 'Strapi', 'KeystoneJS', 'Directus'],
                'Python': ['FastAPI', 'Django', 'Flask', 'Tornado', 'Sanic', 'Pyramid', 
                          'Bottle', 'CherryPy', 'Falcon', 'Starlette'],
                'Java/JVM': ['Spring Boot', 'Micronaut', 'Quarkus', 'Vert.x', 'Play Framework', 
                            'Dropwizard', 'Spark Java', 'Ktor (Kotlin)', 'Akka', 'Helidon'],
                '.NET': ['.NET 8', 'ASP.NET Core', 'Minimal APIs', 'Blazor Server', 'SignalR', 
                        'Orleans', 'MassTransit', 'gRPC', 'Dapr', 'Hot Chocolate'],
                'Go': ['Gin', 'Echo', 'Fiber', 'Chi', 'Gorilla', 'Buffalo', 'Revel', 'Beego', 
                      'Iris', 'Martini'],
                'Rust': ['Actix-web', 'Rocket', 'Axum', 'Warp', 'Tide', 'Tower', 'Hyper', 'Poem',
                        'Salvo', 'Thruster'],
                'Other': ['Ruby on Rails', 'Phoenix (Elixir)', 'Laravel (PHP)', 'Vapor (Swift)', 
                         'Deno', 'Bun', 'Cloudflare Workers', 'Hono', 'Fresh (Deno)', 'Oak (Deno)']
            },
            
            '🗄️ Veritabanı & Veri Depolama': {
                'Relational': ['PostgreSQL', 'MySQL', 'MariaDB', 'SQL Server', 'Oracle', 
                              'SQLite', 'CockroachDB', 'YugabyteDB', 'TiDB', 'AlloyDB'],
                'NoSQL': ['MongoDB', 'DynamoDB', 'Cassandra', 'CouchDB', 'RavenDB', 
                         'ArangoDB', 'OrientDB', 'Amazon DocumentDB', 'Azure Cosmos DB', 'FaunaDB'],
                'Key-Value': ['Redis', 'Memcached', 'KeyDB', 'Hazelcast', 'Aerospike', 
                             'RocksDB', 'LevelDB', 'etcd', 'Amazon ElastiCache', 'DragonflyDB'],
                'Time-Series': ['InfluxDB', 'TimescaleDB', 'Prometheus', 'Graphite', 
                               'OpenTSDB', 'QuestDB', 'VictoriaMetrics', 'Apache Druid', 'ClickHouse', 'TDengine'],
                'Graph': ['Neo4j', 'ArangoDB', 'Amazon Neptune', 'TigerGraph', 'JanusGraph', 
                         'Dgraph', 'OrientDB', 'ArcadeDB', 'RedisGraph', 'Memgraph'],
                'Search': ['Elasticsearch', 'Solr', 'MeiliSearch', 'Typesense', 'Algolia', 
                          'OpenSearch', 'Zinc', 'Sonic', 'Manticore', 'Vespa'],
                'Vector': ['Pinecone', 'Weaviate', 'Qdrant', 'Milvus', 'Chroma', 'Vespa', 
                          'Faiss', 'pgvector', 'Vald', 'Deep Lake']
            },
            
            '🤖 AI/ML & Data Science': {
                'Deep Learning': ['TensorFlow', 'PyTorch', 'JAX', 'Keras', 'MXNet', 
                                 'PaddlePaddle', 'Caffe2', 'ONNX', 'TensorFlow Lite', 'Core ML'],
                'ML Libraries': ['scikit-learn', 'XGBoost', 'LightGBM', 'CatBoost', 
                                'H2O.ai', 'MLlib', 'Prophet', 'PyCaret', 'Auto-sklearn', 'TPOT'],
                'NLP': ['Hugging Face', 'spaCy', 'NLTK', 'Gensim', 'Stanford NLP', 
                       'AllenNLP', 'TextBlob', 'CoreNLP', 'Rasa', 'Haystack'],
                'Computer Vision': ['OpenCV', 'Detectron2', 'YOLO', 'MediaPipe', 'Dlib', 
                                  'SimpleCV', 'Kornia', 'Albumentations', 'MMDetection', 'TorchVision'],
                'LLM/GenAI': ['OpenAI API', 'Anthropic Claude', 'Google Gemini', 'LangChain', 
                             'LlamaIndex', 'Semantic Kernel', 'AutoGen', 'CrewAI', 'Flowise', 'Dify'],
                'MLOps': ['MLflow', 'Kubeflow', 'Weights & Biases', 'Neptune.ai', 'DVC', 
                         'ClearML', 'Metaflow', 'Comet', 'Aim', 'Evidently'],
                'Data Processing': ['Pandas', 'NumPy', 'Apache Spark', 'Dask', 'Ray', 
                                  'Polars', 'Vaex', 'Rapids', 'Modin', 'Apache Flink']
            },
            
            '☁️ Cloud & Infrastructure': {
                'Cloud Providers': ['AWS', 'Azure', 'Google Cloud', 'Oracle Cloud', 'IBM Cloud', 
                                  'Alibaba Cloud', 'Digital Ocean', 'Linode', 'Vultr', 'Hetzner'],
                'Serverless': ['AWS Lambda', 'Azure Functions', 'Google Cloud Functions', 
                              'Vercel', 'Netlify', 'Cloudflare Workers', 'Deno Deploy', 'Railway', 'Render', 'Fly.io'],
                'Container': ['Docker', 'Podman', 'containerd', 'CRI-O', 'LXC/LXD', 
                            'Firecracker', 'gVisor', 'Kata Containers', 'Buildah', 'Kaniko'],
                'Orchestration': ['Kubernetes', 'OpenShift', 'Rancher', 'Docker Swarm', 
                                'Nomad', 'Apache Mesos', 'ECS', 'GKE', 'AKS', 'EKS'],
                'IaC': ['Terraform', 'Ansible', 'Pulumi', 'CloudFormation', 'ARM Templates', 
                       'CDK', 'Crossplane', 'Bicep', 'Saltstack', 'Chef'],
                'CI/CD': ['GitHub Actions', 'GitLab CI', 'Jenkins', 'CircleCI', 'Travis CI', 
                         'Azure DevOps', 'ArgoCD', 'Flux', 'Tekton', 'Drone CI'],
                'Monitoring': ['Prometheus', 'Grafana', 'DataDog', 'New Relic', 'Splunk', 
                             'ELK Stack', 'Jaeger', 'Zipkin', 'AppDynamics', 'Dynatrace']
            },
            
            '⛓️ Blockchain & Web3': {
                'Platforms': ['Ethereum', 'Polygon', 'Binance Smart Chain', 'Solana', 
                             'Avalanche', 'Arbitrum', 'Optimism', 'Cosmos', 'Polkadot', 'Near'],
                'Development': ['Solidity', 'Rust (Solana)', 'Move', 'Cairo', 'Vyper', 
                              'Reach', 'Clarity', 'Michelson', 'Cadence', 'Ink!'],
                'Frameworks': ['Hardhat', 'Truffle', 'Foundry', 'Brownie', 'Anchor', 
                             'Remix IDE', 'OpenZeppelin', 'Dapp Tools', 'Scaffold-ETH', 'Waffle'],
                'Libraries': ['Web3.js', 'Ethers.js', 'Web3.py', 'Wagmi', 'Viem', 
                            'Moralis', 'Alchemy SDK', 'thirdweb', 'QuickNode', 'Infura'],
                'Storage': ['IPFS', 'Arweave', 'Filecoin', 'Storj', 'Sia', 'Swarm', 
                          'Ceramic', 'Gun.js', 'OrbitDB', '3Box'],
                'Tools': ['MetaMask', 'WalletConnect', 'Rainbow Kit', 'Gnosis Safe', 
                        'TheGraph', 'Chainlink', 'Tenderly', 'Etherscan', 'Dune Analytics', 'Nansen']
            },
            
            '📱 Mobile & Cross-Platform': {
                'Native iOS': ['Swift', 'SwiftUI', 'UIKit', 'Core Data', 'CloudKit', 
                             'ARKit', 'Core ML', 'HealthKit', 'Combine', 'RealityKit'],
                'Native Android': ['Kotlin', 'Jetpack Compose', 'Room', 'Hilt', 'Retrofit', 
                                 'Coroutines', 'WorkManager', 'CameraX', 'DataStore', 'Navigation'],
                'Cross-Platform': ['React Native', 'Flutter', 'Ionic', '.NET MAUI', 'NativeScript', 
                                 'Xamarin', 'Capacitor', 'Quasar', 'Framework7', 'Tauri Mobile'],
                'Game Engines': ['Unity', 'Unreal Engine', 'Godot', 'Cocos2d', 'Solar2D', 
                               'Defold', 'GameMaker', 'Construct', 'Phaser', 'PlayCanvas'],
                'Backend Services': ['Firebase', 'Supabase', 'AWS Amplify', 'AppWrite', 
                                   'Parse', 'Back4App', 'Realm', 'OneSignal', 'Pusher', 'PubNub']
            },
            
            '🎮 Gaming & Graphics': {
                'Engines': ['Unity', 'Unreal Engine 5', 'Godot', 'CryEngine', 'Amazon Lumberyard', 
                          'Bevy', 'Panda3D', 'Stride', 'Flax Engine', 'Armory3D'],
                'Graphics APIs': ['OpenGL', 'Vulkan', 'DirectX 12', 'Metal', 'WebGPU', 
                                'WebGL', 'Three.js', 'Babylon.js', 'A-Frame', 'React Three Fiber'],
                'Physics': ['Box2D', 'Bullet', 'PhysX', 'Havok', 'Matter.js', 'Cannon.js', 
                          'Rapier', 'Chipmunk', 'ODE', 'Newton Dynamics'],
                'Audio': ['FMOD', 'Wwise', 'OpenAL', 'Web Audio API', 'Tone.js', 'Howler.js', 
                        'SoLoud', 'BASS', 'Criware', 'Miles'],
                'Networking': ['Mirror', 'Photon', 'Nakama', 'Colyseus', 'Socket.io', 
                             'GameLift', 'PlayFab', 'Steamworks', 'Epic Online Services', 'Netcode']
            },
            
            '🔧 Development Tools': {
                'IDEs': ['VS Code', 'IntelliJ IDEA', 'Visual Studio', 'WebStorm', 'PyCharm', 
                        'Android Studio', 'Xcode', 'Neovim', 'Sublime Text', 'Fleet'],
                'Version Control': ['Git', 'GitHub', 'GitLab', 'Bitbucket', 'Perforce', 
                                  'SVN', 'Mercurial', 'Fossil', 'Azure Repos', 'AWS CodeCommit'],
                'API Tools': ['Postman', 'Insomnia', 'Bruno', 'Thunder Client', 'Hoppscotch', 
                            'REST Client', 'GraphQL Playground', 'Swagger', 'Paw', 'HTTPie'],
                'Design': ['Figma', 'Adobe XD', 'Sketch', 'Framer', 'Penpot', 'Lunacy', 
                         'InVision', 'Principle', 'Zeplin', 'Abstract'],
                'Collaboration': ['Slack', 'Discord', 'Teams', 'Notion', 'Linear', 'Jira', 
                               'Asana', 'Trello', 'Monday', 'ClickUp'],
                'Documentation': ['Docusaurus', 'Gitbook', 'MkDocs', 'Sphinx', 'VuePress', 
                               'Nextra', 'Docsify', 'Read the Docs', 'Mintlify', 'Docz'],
                'Testing': ['Selenium', 'Appium', 'JMeter', 'K6', 'Gatling', 'Locust', 
                          'Artillery', 'SoapUI', 'RestAssured', 'Karate']
            }
        }
        
        # Teknoloji kategorilerini tablo olarak ekle
        for category, subcategories in tech_categories.items():
            self.doc.add_heading(category, level=2)
            
            table = self.doc.add_table(rows=len(subcategories) + 2, cols=2)
            table.style = 'Table Grid'
            
            # Başlık satırı
            table.cell(0, 0).text = 'Alt Kategori'
            table.cell(0, 1).text = 'Teknolojiler (İşaretleyiniz)'
            
            row_idx = 1
            for subcat, techs in subcategories.items():
                table.cell(row_idx, 0).text = subcat
                cell = table.cell(row_idx, 1)
                
                # Teknolojileri checkbox olarak ekle (max 10 göster)
                for tech in techs[:10]:
                    p = cell.add_paragraph()
                    self.add_checkbox(p, tech)
                
                # Her alt kategori için "Diğer" seçeneği
                p = cell.add_paragraph()
                p.add_run('☐ Diğer: _________________')
                
                row_idx += 1
            
            # Genel "Diğer" satırı
            table.cell(row_idx, 0).text = 'Diğer Teknolojiler'
            cell = table.cell(row_idx, 1)
            p = cell.add_paragraph()
            p.add_run('Kullanmayı planladığınız diğer teknolojileri yazınız:\n')
            p.add_run('_' * 60 + '\n' + '_' * 60)
            
            self.doc.add_paragraph()  # Boşluk
            
    def create_project_details_section(self):
        """Proje detayları bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('📋 PROJE DETAYLARI', level=1)
        
        # Problem tanımı
        self.doc.add_heading('Problem Tanımı (5W1H)', level=2)
        
        questions = [
            ('WHAT (Ne)', 'Hangi problemi çözüyorsunuz? Problem tam olarak nedir?', 4),
            ('WHO (Kim)', 'Kimler etkileniyor? Hedef kitle kim? Kullanıcı profili?', 4),
            ('WHERE (Nerede)', 'Problem nerede yaşanıyor? Coğrafi/sektörel kapsam?', 3),
            ('WHEN (Ne zaman)', 'Ne zaman ortaya çıkıyor? Hangi durumlarda?', 3),
            ('WHY (Neden)', 'Neden önemli? Çözülmezse ne olur? Neden şimdi?', 4),
            ('HOW (Nasıl)', 'Nasıl çözmeyi planlıyorsunuz? Yaklaşımınız nedir?', 5)
        ]
        
        for q, desc, lines in questions:
            p = self.doc.add_paragraph()
            run = p.add_run(f'{q}: ')
            run.font.bold = True
            run.font.size = Pt(12)
            p.add_run(f'{desc}\n')
            
            # Yazma alanı
            for _ in range(lines):
                text_box = self.doc.add_paragraph()
                text_box.paragraph_format.left_indent = Inches(0.5)
                text_box.add_run('_' * 90)
            self.doc.add_paragraph()
            
        # Kullanıcı hikayeleri
        self.doc.add_heading('Kullanıcı Hikayeleri (User Stories)', level=2)
        self.doc.add_paragraph('En az 5 kullanıcı hikayesi yazınız (As a... I want... So that...)')
        
        for i in range(1, 6):
            p = self.doc.add_paragraph()
            p.add_run(f'Hikaye {i}:\n')
            p.runs[0].font.bold = True
            
            story_template = self.doc.add_paragraph()
            story_template.paragraph_format.left_indent = Inches(0.5)
            story_template.add_run('As a: ___________________________________________\n')
            story_template.add_run('I want: _________________________________________\n')
            story_template.add_run('So that: ________________________________________')
            self.doc.add_paragraph()
            
    def create_innovation_section(self):
        """İnovasyon ve özgünlük bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('💡 İnovasyon ve Özgünlük', level=1)
        
        self.doc.add_paragraph('Projenizin özgün yanlarını ve yenilikçi özelliklerini açıklayınız:')
        
        # İnovasyon alanları
        innovation_areas = [
            'Teknik İnovasyon (Yeni algoritma, yaklaşım, mimari, teknoloji kullanımı)',
            'İş Modeli İnovasyonu (Yeni gelir modeli, müşteri segmenti, değer önerisi)',
            'Kullanıcı Deneyimi İnovasyonu (Yeni etkileşim yöntemi, arayüz tasarımı)',
            'Süreç İnovasyonu (Yeni çalışma metodolojisi, otomasyon, optimizasyon)',
            'Sosyal İnovasyon (Toplumsal fayda, sürdürülebilirlik, erişilebilirlik)'
        ]
        
        for area in innovation_areas:
            p = self.doc.add_paragraph()
            p.add_run(f'{area}:\n')
            p.runs[0].font.bold = True
            
            for _ in range(3):
                text_line = self.doc.add_paragraph()
                text_line.paragraph_format.left_indent = Inches(0.5)
                text_line.add_run('_' * 85)
            self.doc.add_paragraph()
            
        # Patent potansiyeli
        self.doc.add_heading('Patent/Telif Hakkı Potansiyeli', level=2)
        patent_questions = [
            'Projenizin patentlenebilir yönleri var mı?',
            'Hangi özellikler özgün ve yenilikçi?',
            'Benzer patentler araştırıldı mı?'
        ]
        
        for question in patent_questions:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, question)
            for _ in range(2):
                text_line = self.doc.add_paragraph()
                text_line.paragraph_format.left_indent = Inches(0.5)
                text_line.add_run('_' * 80)
            
    def create_competition_analysis(self):
        """Rekabet analizi bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('🏆 Rekabet Analizi ve Pazar Araştırması', level=1)
        
        # Rakip analiz tablosu
        self.doc.add_heading('Mevcut Çözümler ve Rakipler', level=2)
        
        table = self.doc.add_table(rows=7, cols=6)
        table.style = 'Table Grid'
        
        # Başlıklar
        headers = ['Rakip/Çözüm', 'Güçlü Yanları', 'Zayıf Yanları', 'Fiyat Modeli', 'Pazar Payı', 'Bizim Avantajımız']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        # Boş satırlar
        for i in range(1, 7):
            table.cell(i, 0).text = f'Rakip {i}:'
            
        # SWOT Analizi
        self.doc.add_heading('SWOT Analizi', level=2)
        
        swot_table = self.doc.add_table(rows=3, cols=3)
        swot_table.style = 'Table Grid'
        
        # SWOT başlıkları
        swot_headers = ['', 'Yararlı (Helpful)', 'Zararlı (Harmful)']
        for i, header in enumerate(swot_headers):
            cell = swot_table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        swot_table.cell(1, 0).text = 'İçsel (Internal)'
        swot_table.cell(2, 0).text = 'Dışsal (External)'
        
        swot_table.cell(1, 1).text = 'Güçlü Yanlar (Strengths)\n' + '_' * 30
        swot_table.cell(1, 2).text = 'Zayıf Yanlar (Weaknesses)\n' + '_' * 30
        swot_table.cell(2, 1).text = 'Fırsatlar (Opportunities)\n' + '_' * 30
        swot_table.cell(2, 2).text = 'Tehditler (Threats)\n' + '_' * 30
        
    def create_business_model_section(self):
        """İş modeli ve ticarileşme bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('💰 İŞ MODELİ VE TİCARİLEŞME', level=1)
        
        # Business Model Canvas
        self.doc.add_heading('Business Model Canvas', level=2)
        
        canvas_sections = [
            ('Key Partners', 'Kilit ortaklar, tedarikçiler, iş birlikleri'),
            ('Key Activities', 'Ana faaliyetler, kritik işler'),
            ('Key Resources', 'Kilit kaynaklar (insan, teknoloji, sermaye)'),
            ('Value Propositions', 'Değer önerileri, müşteriye sunulan fayda'),
            ('Customer Relationships', 'Müşteri ilişkileri yönetimi'),
            ('Channels', 'Dağıtım kanalları, müşteriye ulaşma yöntemleri'),
            ('Customer Segments', 'Müşteri segmentleri, hedef kitle'),
            ('Cost Structure', 'Maliyet yapısı, ana gider kalemleri'),
            ('Revenue Streams', 'Gelir akışları, para kazanma yöntemleri')
        ]
        
        for section, desc in canvas_sections:
            p = self.doc.add_paragraph()
            run = p.add_run(f'{section}:')
            run.font.bold = True
            p.add_run(f' ({desc})')
            for _ in range(3):
                self.doc.add_paragraph('_' * 80)
            self.doc.add_paragraph()
            
        # Gelir modeli seçenekleri
        self.doc.add_heading('Gelir Modeli', level=2)
        
        revenue_models = [
            'SaaS (Software as a Service) - Aylık/Yıllık abonelik',
            'Freemium - Temel özellikler ücretsiz, gelişmiş özellikler ücretli',
            'Pay-per-use - Kullanım bazlı ücretlendirme',
            'Lisans satışı - Tek seferlik lisans ücreti',
            'Transaction fee - İşlem başına komisyon',
            'Marketplace - Platform komisyonu',
            'Advertising - Reklam gelirleri',
            'Data monetization - Veri satışı/analitik',
            'Hardware + Software - Donanım ve yazılım paketi',
            'Consulting & Services - Danışmanlık ve özelleştirme',
            'API as a Service - API kullanım ücreti',
            'White-label - Marka lisanslama',
            'Subscription box - Periyodik ürün/hizmet paketi',
            'Crowdfunding - Kitlesel fonlama',
            'Hybrid Model - Karma model',
            'Diğer (Belirtiniz): ______________________'
        ]
        
        for model in revenue_models:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, model)
            
        # Fiyatlama stratejisi
        self.doc.add_heading('Fiyatlama Stratejisi', level=2)
        
        pricing_table = self.doc.add_table(rows=6, cols=5)
        pricing_table.style = 'Table Grid'
        
        # Başlıklar
        headers = ['Plan', 'Hedef Kitle', 'Özellikler', 'Fiyat', 'Kullanıcı Sayısı Hedefi']
        for i, header in enumerate(headers):
            cell = pricing_table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        # Plan örnekleri
        plans = ['Free/Demo', 'Starter', 'Professional', 'Enterprise', 'Custom']
        for i, plan in enumerate(plans, 1):
            pricing_table.cell(i, 0).text = plan
            
    def create_timeline_section(self):
        """Proje zaman planı bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('📅 PROJE ZAMAN PLANI', level=1)
        
        # Sprint planning
        self.doc.add_heading('Sprint Planlaması', level=2)
        
        timeline_table = self.doc.add_table(rows=18, cols=6)
        timeline_table.style = 'Table Grid'
        
        # Başlıklar
        headers = ['Sprint', 'Tarih', 'Aşama', 'Hedefler', 'Deliverables', 'Tamamlanma']
        for i, header in enumerate(headers):
            cell = timeline_table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        # Zaman planı
        timeline = [
            ('Sprint 0', '22 Eyl - 4 Eki', 'Inception', 'Proje kurulumu, takım oluşturma', 'Takım sözleşmesi', '☐'),
            ('Sprint 1', '5 - 18 Eki', 'Research', 'Literatür taraması, pazar araştırması', 'Araştırma raporu', '☐'),
            ('Sprint 2', '19 Eki - 1 Kas', 'Design', 'Sistem tasarımı, PoC geliştirme', 'Tasarım dökümanı', '☐'),
            ('ARA SINAV', '8 - 16 Kas', 'Demo 1', 'Demo (%25), rapor, sunum', 'Ara sınav sunumu', '☐'),
            ('Sprint 3', '17 - 30 Kas', 'Alpha v1', 'Core features geliştirme', 'Alpha v0.1', '☐'),
            ('Sprint 4', '1 - 15 Ara', 'Alpha v2', 'Feature geliştirme', 'Alpha v0.2', '☐'),
            ('Sprint 5', '16 Ara - 2 Oca', 'Testing', 'Test ve iyileştirme', 'Test raporu', '☐'),
            ('FİNAL', '3 - 16 Oca', 'Demo 2', 'Demo (%45), video, iş planı', 'Final sunumu', '☐'),
            ('Sprint 6', '9 - 22 Şub', 'Beta v1', 'Beta geliştirme başlangıcı', 'Beta v0.5', '☐'),
            ('Sprint 7', '23 Şub - 8 Mar', 'User Test', 'Kullanıcı testleri', 'Test feedback', '☐'),
            ('Sprint 8', '9 - 22 Mar', 'Beta v2', 'Beta iyileştirmeler', 'Beta v0.7', '☐'),
            ('ARA SINAV', '28 Mar - 5 Nis', 'Demo 3', 'Demo (%75), kullanıcı feedback', 'Ara sınav sunumu', '☐'),
            ('Sprint 9', '6 - 19 Nis', 'Production', 'Production hazırlık', 'Release candidate', '☐'),
            ('Sprint 10', '20 Nis - 3 May', 'Polish', 'Optimizasyon, güvenlik', 'v0.9', '☐'),
            ('Sprint 11', '4 - 17 May', 'Deploy', 'Deployment, dokümantasyon', 'v1.0', '☐'),
            ('Sprint 12', '18 - 31 May', 'Launch', 'Final hazırlıklar', 'Launch ready', '☐'),
            ('FİNAL', '1 - 14 Haz', 'Demo Final', 'Demo (%100), şirketleşme, sunum', 'Final product', '☐')
        ]
        
        for i, (sprint, date, phase, goal, deliver, status) in enumerate(timeline, 1):
            timeline_table.cell(i, 0).text = sprint
            timeline_table.cell(i, 1).text = date
            timeline_table.cell(i, 2).text = phase
            timeline_table.cell(i, 3).text = goal
            timeline_table.cell(i, 4).text = deliver
            timeline_table.cell(i, 5).text = status
            
        # Milestones
        self.doc.add_heading('Önemli Kilometre Taşları', level=2)
        milestones = [
            'Kasım 2025: Alpha versiyonu tamamlama',
            'Ocak 2026: Beta versiyonu başlatma',
            'Mart 2026: Kullanıcı testleri tamamlama',
            'Mayıs 2026: Production deployment',
            'Haziran 2026: Resmi lansman'
        ]
        
        for milestone in milestones:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, milestone)
            
    def create_success_metrics_section(self):
        """Başarı metrikleri bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('📊 BAŞARI METRİKLERİ VE KPI\'LAR', level=1)
        
        # Metrik kategorileri
        self.doc.add_heading('Teknik Metrikler', level=2)
        tech_metrics = [
            'Kod kalitesi (Code coverage > %70)',
            'Performans (Response time < 500ms)',
            'Güvenlik (OWASP Top 10 uyumlu)',
            'Ölçeklenebilirlik (1000+ eş zamanlı kullanıcı)',
            'Uptime (%99.5 SLA)',
            'Bug yoğunluğu (< 5 bug/KLOC)',
            'API başarı oranı (> %99)',
            'Deployment sıklığı (2+ / hafta)',
            'Mean Time to Recovery (MTTR < 1 saat)',
            'Load time (< 3 saniye)'
        ]
        
        for metric in tech_metrics:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, metric)
            
        # İş metrikleri
        self.doc.add_heading('İş Metrikleri', level=2)
        business_metrics = [
            'Kullanıcı sayısı hedefi: 100+ (6 ay)',
            'Aylık aktif kullanıcı (MAU): 50+',
            'Kullanıcı elde tutma oranı: %60+',
            'Müşteri kazanım maliyeti (CAC): ₺100',
            'Yaşam boyu değer (LTV): ₺1000+',
            'Aylık tekrarlayan gelir (MRR): ₺5000+ (1 yıl)',
            'Churn rate: < %5',
            'NPS skoru: > 50',
            'Conversion rate: > %3',
            'Customer satisfaction: > 4.5/5'
        ]
        
        for metric in business_metrics:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, metric)
            
        # Akademik metrikler
        self.doc.add_heading('Akademik Başarı Kriterleri', level=2)
        academic_metrics = [
            'GitHub\'da 500+ commit',
            'Comprehensive documentation',
            'Academic paper draft hazır',
            'Patent başvurusu yapılmış (opsiyonel)',
            'Yarışmalara başvuru (min. 2)',
            'Demo videoları hazır',
            'Kullanıcı test raporları',
            'Teknik sunum hazır',
            'Poster tasarımı tamamlanmış',
            'Danışman onayı alınmış'
        ]
        
        for metric in academic_metrics:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, metric)
            
        # Özel metrikler için alan
        self.doc.add_heading('Projeye Özel Metrikler', level=2)
        self.doc.add_paragraph('Projenize özgü başarı kriterlerini belirtiniz:')
        
        for _ in range(5):
            p = self.doc.add_paragraph()
            p.add_run('• ')
            p.add_run('_' * 80)
            
    def create_risk_analysis_section(self):
        """Risk analizi bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('⚠️ RİSK ANALİZİ VE YÖNETİMİ', level=1)
        
        # Risk matrisi
        self.doc.add_heading('Risk Değerlendirme Matrisi', level=2)
        
        risk_table = self.doc.add_table(rows=13, cols=7)
        risk_table.style = 'Table Grid'
        
        # Başlıklar
        headers = ['Risk Kategorisi', 'Risk Açıklaması', 'Olasılık\n(1-5)', 'Etki\n(1-5)', 
                  'Risk Skoru', 'Azaltma Stratejisi', 'Sorumlu']
        for i, header in enumerate(headers):
            cell = risk_table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        # Risk kategorileri
        risk_categories = [
            'Teknik Risk',
            'Pazar Riski',
            'Finansal Risk',
            'Takım Riski',
            'Zaman Riski',
            'Yasal Risk',
            'Güvenlik Riski',
            'Operasyonel Risk',
            'Rekabet Riski',
            'Teknoloji Riski',
            'Müşteri Riski',
            'Diğer Riskler'
        ]
        
        for i, category in enumerate(risk_categories, 1):
            risk_table.cell(i, 0).text = category
            
        # Risk mitigation plan
        self.doc.add_heading('Risk Azaltma Planı', level=2)
        self.doc.add_paragraph('Kritik riskler için detaylı aksiyon planı:')
        
        for i in range(1, 4):
            p = self.doc.add_paragraph()
            p.add_run(f'Risk {i}:\n')
            p.runs[0].font.bold = True
            
            for _ in range(3):
                text_line = self.doc.add_paragraph()
                text_line.paragraph_format.left_indent = Inches(0.5)
                text_line.add_run('_' * 80)
            self.doc.add_paragraph()
            
    def create_resources_section(self):
        """Kaynaklar ve ihtiyaçlar bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('🔧 KAYNAKLAR VE İHTİYAÇLAR', level=1)
        
        # Bütçe tablosu
        self.doc.add_heading('Bütçe Planlaması', level=2)
        
        budget_table = self.doc.add_table(rows=15, cols=5)
        budget_table.style = 'Table Grid'
        
        # Başlıklar
        headers = ['Kategori', 'Açıklama', 'Miktar', 'Tahmini Maliyet (₺)', 'Notlar']
        for i, header in enumerate(headers):
            cell = budget_table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        # Bütçe kalemleri
        budget_items = [
            'Hardware/Donanım',
            'Cloud/Hosting',
            'Yazılım Lisansları',
            'API Kullanımları',
            'Domain/SSL',
            'Marketing/Tanıtım',
            'Test Kullanıcı Ödemeleri',
            'Yarışma Başvuruları',
            'Patent/Telif',
            'Eğitim/Sertifika',
            'Danışmanlık',
            'Prototip Malzemeleri',
            'Diğer',
            'TOPLAM'
        ]
        
        for i, item in enumerate(budget_items, 1):
            budget_table.cell(i, 0).text = item
            if item == 'TOPLAM':
                for p in budget_table.cell(i, 0).paragraphs:
                    p.runs[0].font.bold = True
                    
        # İnsan kaynakları
        self.doc.add_heading('İnsan Kaynakları', level=2)
        
        team_resources = [
            'Takım üyeleri yetkinlikleri yeterli mi?',
            'Ek teknik destek gerekiyor mu?',
            'Mentor/danışman desteği var mı?',
            'Domain expert erişimi var mı?'
        ]
        
        for resource in team_resources:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, resource)
            
    def create_testing_section(self):
        """Test stratejisi bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('🧪 TEST STRATEJİSİ', level=1)
        
        # Test türleri
        self.doc.add_heading('Test Türleri', level=2)
        
        test_types = [
            'Unit Testing (Birim testleri)',
            'Integration Testing (Entegrasyon testleri)',
            'System Testing (Sistem testleri)',
            'User Acceptance Testing (Kullanıcı kabul testleri)',
            'Performance Testing (Performans testleri)',
            'Security Testing (Güvenlik testleri)',
            'Usability Testing (Kullanılabilirlik testleri)',
            'Compatibility Testing (Uyumluluk testleri)',
            'Regression Testing (Regresyon testleri)',
            'A/B Testing'
        ]
        
        for test_type in test_types:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, test_type)
            
        # Test planı
        self.doc.add_heading('Test Planı', level=2)
        
        test_plan_table = self.doc.add_table(rows=6, cols=5)
        test_plan_table.style = 'Table Grid'
        
        headers = ['Test Aşaması', 'Test Türü', 'Araçlar', 'Metrikler', 'Zaman']
        for i, header in enumerate(headers):
            cell = test_plan_table.cell(0, i)
            cell.text = header
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                
        test_phases = ['Development', 'Alpha', 'Beta', 'Pre-Production', 'Production']
        for i, phase in enumerate(test_phases, 1):
            test_plan_table.cell(i, 0).text = phase
            
    def create_marketing_section(self):
        """Pazarlama ve tanıtım stratejisi"""
        self.doc.add_page_break()
        self.doc.add_heading('📢 PAZARLAMA VE TANITIM STRATEJİSİ', level=1)
        
        # Pazarlama kanalları
        self.doc.add_heading('Pazarlama Kanalları', level=2)
        
        marketing_channels = [
            'Social Media (LinkedIn, Twitter, Instagram)',
            'Content Marketing (Blog, Medium)',
            'SEO & SEM',
            'Email Marketing',
            'Influencer Marketing',
            'Community Building (Discord, Slack)',
            'Product Hunt Launch',
            'Hacker News',
            'Reddit Communities',
            'YouTube (Demo videos)',
            'Podcast görünürlüğü',
            'Webinar & Workshop',
            'Networking Events',
            'University Partnerships',
            'Tech Conferences'
        ]
        
        for channel in marketing_channels:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, channel)
            
        # Go-to-market stratejisi
        self.doc.add_heading('Go-to-Market Stratejisi', level=2)
        
        gtm_phases = [
            ('Soft Launch', 'Beta kullanıcıları ile test'),
            ('Product Hunt', 'Product Hunt\'ta lansman'),
            ('PR Campaign', 'Basın bültenleri ve haber siteleri'),
            ('Content Push', 'Blog yazıları ve case study\'ler'),
            ('Partnership', 'Stratejik ortaklıklar'),
            ('Scale', 'Paid marketing ve ölçeklendirme')
        ]
        
        for phase, desc in gtm_phases:
            p = self.doc.add_paragraph()
            p.add_run(f'{phase}: ')
            p.runs[0].font.bold = True
            p.add_run(desc)
            
    def create_appendix_section(self):
        """Ekler bölümü"""
        self.doc.add_page_break()
        self.doc.add_heading('📎 EKLER', level=1)
        
        # Yarışma listesi
        self.doc.add_heading('EK-1: Hedef Yarışmalar', level=2)
        
        competitions = [
            'TEKNOFEST (Çeşitli kategoriler)',
            'TÜBİTAK 2242 Üniversite Öğrencileri Araştırma Proje Yarışması',
            'Google Solution Challenge',
            'Microsoft Imagine Cup',
            'NASA Space Apps Challenge',
            'Hackathon\'lar',
            'Startup Weekend',
            'Big Bang Startup Challenge',
            'Garanti BBVA Partners Accelerator',
            'İTÜ Çekirdek Big Bang',
            'Endeavor Turkey',
            'Workup Girişimcilik Programı'
        ]
        
        for comp in competitions:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, comp)
            
        # Haftalık toplantı şablonu
        self.doc.add_heading('EK-2: Haftalık Toplantı Şablonu', level=2)
        
        meeting_template = """
        Tarih: ___________
        Katılımcılar: ___________
        
        📝 Tamamlananlar:
        • ___________
        • ___________
        
        🔄 Devam Edenler:
        • ___________
        • ___________
        
        🚫 Blokajlar:
        • ___________
        
        📌 Gelecek Hafta:
        • ___________
        • ___________
        
        Action Items:
        □ @isim: görev (deadline)
        □ @isim: görev (deadline)
        """
        
        self.doc.add_paragraph(meeting_template)
        
        # İletişim bilgileri
        self.doc.add_heading('EK-3: İletişim Bilgileri', level=2)
        
        contact_info = self.doc.add_table(rows=8, cols=2)
        contact_info.style = 'Table Grid'
        
        contacts = [
            ('Danışman', 'Dr. Uğur CORUH'),
            ('E-posta', 'ugur.coruh@erdogan.edu.tr'),
            ('Ofis', '+90 (464) 223 75 18 / 1246'),
            ('Adres', 'Zihni Derin Yerleşkesi, Fener Mah. 53100 Rize'),
            ('GitHub Org', 'github.com/rteu-ceng'),
            ('Toplantı Zamanı', 'Her Pazartesi 14:00'),
            ('Discord', 'RTEÜ CS Server'),
            ('Web', 'bilgisayar.erdogan.edu.tr')
        ]
        
        for i, (label, value) in enumerate(contacts):
            contact_info.cell(i, 0).text = label
            contact_info.cell(i, 1).text = value
            
        # Önemli linkler
        self.doc.add_heading('EK-4: Faydalı Kaynaklar ve Linkler', level=2)
        
        links = [
            'RTEÜ Bilgisayar Mühendisliği: https://bilgisayar.erdogan.edu.tr',
            'Akademik Takvim: https://www.erdogan.edu.tr/tr/akademik-takvim',
            'TEKNOFEST: https://www.teknofest.org',
            'TÜBİTAK 2242: https://tubitak.gov.tr',
            'GitHub Student Pack: https://education.github.com',
            'Microsoft Azure for Students: https://azure.microsoft.com/free/students',
            'AWS Educate: https://aws.amazon.com/education',
            'Google Cloud for Students: https://cloud.google.com/edu',
            'JetBrains Student License: https://www.jetbrains.com/student',
            'Figma Education: https://www.figma.com/education',
            'Notion for Education: https://www.notion.so/product/notion-for-education',
            'Canva for Education: https://www.canva.com/education'
        ]
        
        for link in links:
            p = self.doc.add_paragraph()
            p.add_run('• ' + link)
            
        # Notlar alanı
        self.doc.add_heading('EK-5: Notlar ve Düşünceler', level=2)
        for _ in range(15):
            p = self.doc.add_paragraph()
            p.add_run('_' * 100)
            
    def add_final_checklist(self):
        """Final kontrol listesi"""
        self.doc.add_page_break()
        self.doc.add_heading('✅ FİNAL KONTROL LİSTESİ', level=1)
        
        self.doc.add_heading('Proje Başlangıcı', level=2)
        checklist_start = [
            'Takım sözleşmesi imzalandı',
            'GitHub repository oluşturuldu',
            'README.md hazırlandı',
            'Proje klasör yapısı oluşturuldu',
            'Development environment kuruldu',
            '.gitignore ve .env.example hazır',
            'CI/CD pipeline kuruldu',
            'İletişim kanalları kuruldu (Discord/Slack)',
            'Proje yönetim aracı seçildi (Jira/Trello/Notion)',
            'Haftalık toplantı zamanı belirlendi'
        ]
        
        for item in checklist_start:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, item)
            
        self.doc.add_heading('Planlama ve Tasarım', level=2)
        checklist_planning = [
            'Proje başlığı ve domain belirlendi',
            'Problem tanımı netleştirildi',
            'Kullanıcı hikayeleri yazıldı',
            'Teknoloji stack\'i kararlaştırıldı',
            'Sistem mimarisi tasarlandı',
            'Veritabanı şeması hazırlandı',
            'API tasarımı tamamlandı',
            'UI/UX mockup\'ları hazır',
            'Sprint planı hazırlandı',
            'Risk analizi yapıldı'
        ]
        
        for item in checklist_planning:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, item)
            
        self.doc.add_heading('Geliştirme', level=2)
        checklist_dev = [
            'Kod standartları belirlendi',
            'Code review süreci tanımlandı',
            'Test stratejisi belirlendi',
            'Güvenlik kontrolleri yapıldı',
            'Performance optimizasyonu yapıldı',
            'Dokümantasyon güncel',
            'Deployment pipeline hazır',
            'Monitoring ve logging kuruldu',
            'Backup stratejisi belirlendi',
            'Disaster recovery planı hazır'
        ]
        
        for item in checklist_dev:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, item)
            
        self.doc.add_heading('Teslim ve Sunum', level=2)
        checklist_delivery = [
            'İlk literatür taraması yapıldı',
            'Rakip analizi tamamlandı',
            'İş modeli belirlendi',
            'Bütçe planlaması tamamlandı',
            'Patent araştırması yapıldı',
            'Demo videoları hazırlandı',
            'Sunum dosyası hazırlandı',
            'Poster tasarımı tamamlandı',
            'Yarışma başvuruları yapıldı',
            'Danışman onayı alındı'
        ]
        
        for item in checklist_delivery:
            p = self.doc.add_paragraph()
            self.add_checkbox(p, item)
            
    def generate_template(self, filename='RTEU_Bitirme_Tezi__Sablon.docx'):
        """Şablonu oluştur ve kaydet"""
        print("📝  şablon oluşturuluyor...")
        
        # Üst/Alt bilgi
        self.add_header_footer()
        
        # Ana bölümler
        self.create_cover_page()
        self.create_team_info_section()
        self.create_comprehensive_domains()
        self.create_comprehensive_tech_stack()
        self.create_project_details_section()
        self.create_innovation_section()
        self.create_competition_analysis()
        self.create_business_model_section()
        self.create_timeline_section()
        self.create_success_metrics_section()
        self.create_risk_analysis_section()
        self.create_resources_section()
        self.create_testing_section()
        self.create_marketing_section()
        self.create_appendix_section()
        self.add_final_checklist()
        
        # İmza bölümü
        self.doc.add_page_break()
        self.doc.add_heading('İMZALAR', level=1)
        
        signatures = [
            'Takım Lideri',
            'Teknik Lider',
            'Takım Üyesi 3',
            'Danışman (Dr. Uğur CORUH)'
        ]
        
        for sig in signatures:
            p = self.doc.add_paragraph()
            p.add_run(f'{sig}:')
            p.add_run('\t' * 3)
            p.add_run('_' * 30)
            p.add_run('\t')
            p.add_run('Tarih: ___/___/2025')
            self.doc.add_paragraph()
            self.doc.add_paragraph()
            
        # Motivasyon
        self.doc.add_paragraph()
        quote = self.doc.add_paragraph()
        quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        quote.add_run('"The best way to predict the future is to invent it."\n')
        quote.add_run('- Alan Kay\n\n')
        quote.add_run('"Move fast and break things."\n')
        quote.add_run('- Mark Zuckerberg\n\n')
        quote.add_run('Başarılar! 🚀')
        for i in [1, 3]:
            quote.runs[i].font.italic = True
            quote.runs[i].font.size = Pt(10)
        
        # Belgeyi kaydet
        self.doc.save(filename)
        print(f"✅ Şablon başarıyla oluşturuldu: {filename}")
        print(f"📄 Dosya boyutu: {os.path.getsize(filename) / 1024:.2f} KB")
        
        return filename


class ThesisProjectManager:
    """Proje yönetim araçları"""
    
    def __init__(self):
        self.projects = []
        
    def create_project_folder_structure(self, project_name):
        """Proje klasör yapısını oluştur"""
        base_path = f"RTEU_BTZ_{project_name}"
        
        folders = [
            f"{base_path}/01_Dokümantasyon",
            f"{base_path}/01_Dokümantasyon/Raporlar",
            f"{base_path}/01_Dokümantasyon/Sunumlar",
            f"{base_path}/01_Dokümantasyon/Toplantı_Notları",
            f"{base_path}/02_Kod",
            f"{base_path}/02_Kod/Frontend",
            f"{base_path}/02_Kod/Backend",
            f"{base_path}/02_Kod/Database",
            f"{base_path}/02_Kod/Mobile",
            f"{base_path}/02_Kod/Tests",
            f"{base_path}/02_Kod/Scripts",
            f"{base_path}/03_Tasarım",
            f"{base_path}/03_Tasarım/UI_UX",
            f"{base_path}/03_Tasarım/Mockups",
            f"{base_path}/03_Tasarım/Prototypes",
            f"{base_path}/03_Tasarım/Assets",
            f"{base_path}/04_Araştırma",
            f"{base_path}/04_Araştırma/Literatür",
            f"{base_path}/04_Araştırma/Pazar_Analizi",
            f"{base_path}/04_Araştırma/Rakip_Analizi",
            f"{base_path}/05_Toplantılar",
            f"{base_path}/05_Toplantılar/Haftalık",
            f"{base_path}/05_Toplantılar/Sprint_Review",
            f"{base_path}/06_Yarışmalar",
            f"{base_path}/06_Yarışmalar/TEKNOFEST",
            f"{base_path}/06_Yarışmalar/TUBITAK",
            f"{base_path}/07_Medya",
            f"{base_path}/07_Medya/Videolar",
            f"{base_path}/07_Medya/Posterler",
            f"{base_path}/07_Medya/Screenshots",
            f"{base_path}/08_Deployment",
            f"{base_path}/08_Deployment/Docker",
            f"{base_path}/08_Deployment/Kubernetes",
            f"{base_path}/08_Deployment/CI_CD",
            f"{base_path}/09_Test",
            f"{base_path}/09_Test/Unit_Tests",
            f"{base_path}/09_Test/Integration_Tests",
            f"{base_path}/09_Test/User_Tests",
            f"{base_path}/10_Backup"
        ]
        
        for folder in folders:
            os.makedirs(folder, exist_ok=True)
            
        # README dosyaları oluştur
        readme_content = f"""# {project_name} - Bitirme Tezi Projesi

## 📁 Klasör Yapısı

- **01_Dokümantasyon**: Tüm proje dökümanları
- **02_Kod**: Kaynak kodları
- **03_Tasarım**: UI/UX tasarımları
- **04_Araştırma**: Literatür ve pazar araştırması
- **05_Toplantılar**: Toplantı notları
- **06_Yarışmalar**: Yarışma başvuruları
- **07_Medya**: Videolar ve görseller
- **08_Deployment**: Deployment konfigürasyonları
- **09_Test**: Test dosyaları
- **10_Backup**: Yedekler

## 🚀 Hızlı Başlangıç

1. Gerekli bağımlılıkları yükleyin
2. Konfigürasyon dosyasını ayarlayın
3. Projeyi çalıştırın

## 💻 Teknoloji Stack

[Teknolojiler buraya eklenecek]

## 👥 Takım

- Takım Lideri: [İsim]
- Teknik Lider: [İsim]
- Danışman: Dr. Uğur CORUH

## 📅 Önemli Tarihler

- Ara Sınav Demo: 8-16 Kasım 2025
- Final Demo: 3-16 Ocak 2026
- Beta Release: Mart 2026
- Production: Haziran 2026

## 📝 Lisans

Bu proje RTEÜ Bilgisayar Mühendisliği Bölümü bitirme tezi kapsamında geliştirilmektedir.

## 🤝 Katkıda Bulunma

1. Fork yapın
2. Feature branch oluşturun (`git checkout -b feature/AmazingFeature`)
3. Commit yapın (`git commit -m 'Add some AmazingFeature'`)
4. Push yapın (`git push origin feature/AmazingFeature`)
5. Pull Request açın

## 📧 İletişim

- Dr. Uğur CORUH - ugur.coruh@erdogan.edu.tr
- Proje GitHub: [URL]
"""
        
        with open(f"{base_path}/README.md", 'w', encoding='utf-8') as f:
            f.write(readme_content)
            
        # .gitignore oluştur
        gitignore_content = """# Python
*.pyc
__pycache__/
venv/
env/
.env
.venv
*.egg-info/
dist/
build/

# IDE
.vscode/
.idea/
*.swp
*.swo
.DS_Store

# Node
node_modules/
npm-debug.log
yarn-error.log
package-lock.json
yarn.lock

# OS
.DS_Store
Thumbs.db
desktop.ini

# Project specific
*.log
*.tmp
*.cache
.coverage
htmlcov/
.pytest_cache/

# Build
dist/
build/
*.exe
*.dll
*.so
*.dylib

# Database
*.db
*.sqlite
*.sqlite3

# Secrets
secrets.json
config.json
*.pem
*.key

# Backup
*.bak
*.backup
"""
        
        with open(f"{base_path}/.gitignore", 'w', encoding='utf-8') as f:
            f.write(gitignore_content)
            
        # requirements.txt oluştur
        requirements_content = """# Core dependencies
python-docx==0.8.11
pandas==2.0.0
numpy==1.24.0
requests==2.31.0

# Web framework
fastapi==0.100.0
uvicorn==0.23.0

# Database
sqlalchemy==2.0.0
psycopg2-binary==2.9.0

# Testing
pytest==7.4.0
pytest-cov==4.1.0

# Development
black==23.0.0
flake8==6.0.0
pre-commit==3.3.0
"""
        
        with open(f"{base_path}/requirements.txt", 'w', encoding='utf-8') as f:
            f.write(requirements_content)
            
        print(f"✅ Proje klasör yapısı oluşturuldu: {base_path}")
        print(f"📁 Toplam {len(folders)} klasör oluşturuldu")
        print(f"📄 README.md, .gitignore ve requirements.txt dosyaları hazır")
        
    def create_multiple_templates(self, team_count=5):
        """Birden fazla takım için şablon oluştur"""
        for i in range(1, team_count + 1):
            generator = RTEUThesisTemplate()
            filename = f"RTEU_Bitirme_Tezi_Sablon_Takim_{i}.docx"
            generator.generate_template(filename)
            print(f"📄 Takım {i} şablonu oluşturuldu")


def main():
    """Ana fonksiyon"""
    print("=" * 70)
    print("🎓 RTEÜ Bilgisayar Mühendisliği")
    print("📚 Bitirme Tezi  Şablon Oluşturucu v5.0")
    print("=" * 70)
    
    #  şablon oluştur
    generator = RTEUThesisTemplate()
    template_file = generator.generate_template()
    
    print("\n📊 Şablon İçeriği:")
    print("✔ 20+ Domain kategorisi")
    print("✔ 300+ Alt kategori")
    print("✔ 400+ Teknoloji seçeneği")
    print("✔ Kapsamlı proje yönetimi bölümleri")
    print("✔ Detaylı risk analizi ve metrikler")
    print("✔ Pazarlama ve test stratejileri")
    print("✔ Yarışma listesi ve kaynaklar")
    print("✔ Final kontrol listesi")
    
    print("\n📋 Ek işlemler:")
    
    # Proje yöneticisi
    manager = ThesisProjectManager()
    
    # Örnek proje klasörü oluştur
    create_folder = input("📁 Proje klasör yapısı oluşturulsun mu? (E/H): ")
    if create_folder.upper() == 'E':
        project_name = input("Proje adı (boşluksuz): ")
        manager.create_project_folder_structure(project_name)
    
    # Birden fazla şablon
    create_multiple = input("📄 Birden fazla takım şablonu oluşturulsun mu? (E/H): ")
    if create_multiple.upper() == 'E':
        count = int(input("Kaç takım için? "))
        manager.create_multiple_templates(count)
    
    print("\n✨ Tüm işlemler tamamlandı!")
    print("📧 Sorularınız için: ugur.coruh@erdogan.edu.tr")
    print("🚀 Başarılar dilerim!")
    print("\n💡 İpucu: Oluşturulan şablonu dikkatlice doldurun ve")
    print("   danışmanınızla düzenli olarak görüşmeyi unutmayın!")

if __name__ == "__main__":
    main()